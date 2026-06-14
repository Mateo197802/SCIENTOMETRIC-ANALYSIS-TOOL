"""Build the complete ML4Africa manuscript as an editable Word document."""

from __future__ import annotations

import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BASE_DIR = Path(__file__).resolve().parents[2]
MAIN_SOURCE = BASE_DIR / "manuscript" / "text" / "manuscript.md"
SUPPLEMENT_SOURCE = BASE_DIR / "manuscript" / "text" / "supplement.md"
REFERENCE_SOURCE = (
    BASE_DIR / "manuscript" / "references" / "references_vancouver.md"
)
OUTPUT_PATH = (
    BASE_DIR / "output" / "doc" / "ML4Africa_scientometric_manuscript.docx"
)

APPROVED_TITLE = (
    "Geographic Patterns of Scientific Leadership and Bibliometric Impact "
    "in the ML4Africa Research Corpus"
)

NAVY = "17324D"
TEAL = "0B7A75"
LIGHT_TEAL = "EAF7F5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
YELLOW = "FFF4CC"
WHITE = "FFFFFF"

FIGURES = {
    "Figure 1": BASE_DIR / "manuscript" / "figures" / "figure_1_pipeline.png",
    "Figure 2": (
        BASE_DIR
        / "manuscript"
        / "figures"
        / "figure_2_corpus_collaboration.png"
    ),
    "Figure 3": (
        BASE_DIR / "manuscript" / "figures" / "figure_3_mixed_leadership.png"
    ),
    "Figure 4": (
        BASE_DIR / "manuscript" / "figures" / "figure_4_regional_hindex.png"
    ),
    "Figure 5": (
        BASE_DIR / "manuscript" / "figures" / "figure_5_country_hindex.png"
    ),
}

TABLES = {
    "Table 1": (
        BASE_DIR / "manuscript" / "tables" / "table_1_corpus_characteristics.csv"
    ),
    "Table 2": (
        BASE_DIR
        / "manuscript"
        / "tables"
        / "table_2_collaboration_leadership.csv"
    ),
    "Table 3": (
        BASE_DIR / "manuscript" / "tables" / "table_3_bibliometric_impact.csv"
    ),
}

SUPPLEMENTARY_TABLES = {
    "S1": (
        BASE_DIR
        / "manuscript"
        / "tables"
        / "supplementary_variable_dictionary.csv"
    ),
    "S2": BASE_DIR / "data" / "analysis" / "field_coverage.csv",
    "S3": (
        BASE_DIR / "manuscript" / "tables" / "supplementary_publication_years.csv"
    ),
    "S4": (
        BASE_DIR
        / "manuscript"
        / "tables"
        / "supplementary_country_participation.csv"
    ),
    "S5": (
        BASE_DIR
        / "manuscript"
        / "tables"
        / "supplementary_country_leadership.csv"
    ),
    "S6": (
        BASE_DIR
        / "manuscript"
        / "tables"
        / "supplementary_country_impact.csv"
    ),
    "S7": (
        BASE_DIR
        / "manuscript"
        / "tables"
        / "supplementary_gender_role_summary.csv"
    ),
    "S8": (
        BASE_DIR
        / "manuscript"
        / "tables"
        / "supplementary_profile_summary.csv"
    ),
}

PLACEHOLDER_PATTERN = re.compile(r"^\{\{([^{}]+)\}\}$")
INLINE_PATTERN = re.compile(r"(https?://[^\s]+|\*\*.+?\*\*|`.+?`)")


def _set_cell_shading(cell, fill: str) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = cell_properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        cell_properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def _set_cell_margins(cell, top=60, start=70, bottom=60, end=70) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for margin_name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = margins.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), TEAL)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([color, underline])
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_inline_content(paragraph, text: str) -> None:
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        token = match.group(0)
        if token.startswith("http"):
            url = token.rstrip(".,;)")
            suffix = token[len(url) :]
            _add_hyperlink(paragraph, url, url)
            if suffix:
                paragraph.add_run(suffix)
        elif token.startswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _style_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.widow_control = True

    title = document.styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(16)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    title.paragraph_format.space_after = Pt(12)

    for style_name, size, color in (
        ("Heading 1", 14, NAVY),
        ("Heading 2", 12, TEAL),
        ("Heading 3", 11, NAVY),
    ):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)

    caption = document.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(MID_GRAY)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_together = True

    if "References" not in document.styles:
        style = document.styles.add_style("References", WD_STYLE_TYPE.PARAGRAPH)
    reference_style = document.styles["References"]
    reference_style.font.name = "Times New Roman"
    reference_style.font.size = Pt(9)
    reference_style.paragraph_format.left_indent = Inches(0.25)
    reference_style.paragraph_format.first_line_indent = Inches(-0.25)
    reference_style.paragraph_format.space_after = Pt(4)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    _add_page_number(footer)

    header = section.header.paragraphs[0]
    header.text = "ML4Africa scientometric manuscript"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MID_GRAY)


def _format_header(value: str) -> str:
    result = value.replace("_", " ").title()
    replacements = {
        "Doi": "DOI",
        "Iqr": "IQR",
        "Hindex": "H-index",
        "Orcid": "ORCID",
        "Pmid": "PMID",
        "Llm": "LLM",
        "Oa": "OA",
        "Ss": "SS",
    }
    for old, new in replacements.items():
        result = result.replace(old, new)
    return result


def _format_value(value) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, float):
        return f"{value:.1f}"
    return str(value)


def _add_table(document: Document, path: Path, supplementary: bool) -> None:
    if not path.exists():
        raise FileNotFoundError(path)
    data = pd.read_csv(path)
    table = document.add_table(rows=1, cols=len(data.columns))
    table.style = "Table Grid"
    table.autofit = True
    header_cells = table.rows[0].cells
    _set_repeat_table_header(table.rows[0])
    font_size = 6.5 if supplementary else 8

    for index, column in enumerate(data.columns):
        cell = header_cells[index]
        cell.text = _format_header(str(column))
        _set_cell_shading(cell, NAVY)
        _set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(font_size)
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(WHITE)

    for row_index, values in enumerate(data.itertuples(index=False, name=None), start=1):
        cells = table.add_row().cells
        for column_index, value in enumerate(values):
            cell = cells[column_index]
            cell.text = _format_value(value)
            if row_index % 2 == 0:
                _set_cell_shading(cell, LIGHT_TEAL)
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = RGBColor.from_string(NAVY)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_figure(document: Document, path: Path) -> None:
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    paragraph.add_run().add_picture(str(path), width=Inches(6.45))


def _add_references(document: Document) -> None:
    references = [
        line.strip()
        for line in REFERENCE_SOURCE.read_text(encoding="utf-8").splitlines()
        if line.strip()
    ]
    for reference in references:
        paragraph = document.add_paragraph(style="References")
        _add_inline_content(paragraph, reference)


def _make_landscape_section(document: Document) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header.is_linked_to_previous = True
    section.footer.is_linked_to_previous = True


def _add_heading(document: Document, text: str, level: int, supplement: bool) -> None:
    if level == 1 and not supplement and not document.paragraphs:
        paragraph = document.add_paragraph(text, style="Title")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return
    style_level = min(max(level, 1), 3)
    document.add_heading(text, level=style_level)


def _add_text_paragraph(document: Document, text: str, style: str | None = None):
    paragraph = document.add_paragraph(style=style)
    if text.startswith("[Editorial completion required:"):
        _set_cell_shading_proxy(paragraph, YELLOW)
    _add_inline_content(paragraph, text)
    return paragraph


def _set_cell_shading_proxy(paragraph, fill: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _handle_placeholder(
    document: Document,
    placeholder: str,
    supplement: bool,
) -> None:
    if placeholder.startswith("FIGURE:"):
        key = f"Figure {placeholder.split(':', 1)[1]}"
        if key not in FIGURES:
            raise ValueError(f"Unknown figure placeholder: {placeholder}")
        _add_figure(document, FIGURES[key])
        return
    if placeholder.startswith("TABLE:"):
        key = f"Table {placeholder.split(':', 1)[1]}"
        if key not in TABLES:
            raise ValueError(f"Unknown table placeholder: {placeholder}")
        _add_table(document, TABLES[key], supplementary=False)
        return
    if placeholder.startswith("SUPPLEMENTARY_TABLE:"):
        key = placeholder.split(":", 1)[1]
        if key not in SUPPLEMENTARY_TABLES:
            raise ValueError(f"Unknown supplementary table placeholder: {placeholder}")
        _add_table(document, SUPPLEMENTARY_TABLES[key], supplementary=True)
        return
    if placeholder == "REFERENCES":
        _add_references(document)
        return
    if placeholder == "SUPPLEMENT":
        _make_landscape_section(document)
        _render_markdown(
            document,
            SUPPLEMENT_SOURCE.read_text(encoding="utf-8"),
            supplement=True,
        )
        return
    raise ValueError(f"Unknown manuscript placeholder: {placeholder}")


def _render_markdown(document: Document, text: str, supplement: bool) -> None:
    lines = text.splitlines()
    paragraph_buffer: list[str] = []

    def flush_buffer() -> None:
        if paragraph_buffer:
            joined = " ".join(line.strip() for line in paragraph_buffer)
            _add_text_paragraph(document, joined)
            paragraph_buffer.clear()

    for raw_line in lines:
        line = raw_line.rstrip()
        if not line.strip():
            flush_buffer()
            continue

        placeholder_match = PLACEHOLDER_PATTERN.match(line.strip())
        if placeholder_match:
            flush_buffer()
            _handle_placeholder(
                document,
                placeholder_match.group(1),
                supplement=supplement,
            )
            continue

        if "{{" in line or "}}" in line:
            raise ValueError(f"Malformed manuscript placeholder: {line}")

        heading_match = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading_match:
            flush_buffer()
            heading_text = heading_match.group(2).strip()
            if heading_text == "References":
                document.add_page_break()
            _add_heading(
                document,
                heading_text,
                len(heading_match.group(1)),
                supplement=supplement,
            )
            continue

        if line.startswith("- "):
            flush_buffer()
            _add_text_paragraph(document, line[2:].strip(), style="List Bullet")
            continue

        if re.match(r"^\d+\.\s+", line):
            flush_buffer()
            _add_text_paragraph(
                document,
                re.sub(r"^\d+\.\s+", "", line),
                style="List Number",
            )
            continue

        if line.startswith("**Figure") or line.startswith("**Table"):
            flush_buffer()
            _add_text_paragraph(document, line, style="Caption")
            continue
        if line.startswith("**Supplementary Table"):
            flush_buffer()
            _add_text_paragraph(document, line, style="Caption")
            continue

        paragraph_buffer.append(line)

    flush_buffer()


def build_document(output_path: Path = OUTPUT_PATH) -> Path:
    """Build and save the complete manuscript Word file."""
    document = Document()
    _style_document(document)
    document.core_properties.title = APPROVED_TITLE
    document.core_properties.subject = "ML4Africa scientometric analysis"
    document.core_properties.author = (
        "ML4Africa research team - author list pending editorial confirmation"
    )
    document.core_properties.keywords = (
        "scientometrics; Africa; machine learning for health; authorship"
    )

    source = MAIN_SOURCE.read_text(encoding="utf-8")
    _render_markdown(document, source, supplement=False)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def main() -> None:
    output = build_document()
    print(f"Word manuscript built: {output}")


if __name__ == "__main__":
    main()
